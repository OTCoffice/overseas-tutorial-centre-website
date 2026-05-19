from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
DOCX_OUT = ASSETS / "OTC_University_Representative_Capability_Statement_2026.docx"
PDF_OUT = ASSETS / "OTC_University_Representative_Capability_Statement_2026.pdf"

NAVY = RGBColor(6, 26, 47)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(184, 138, 59)
MUTED = RGBColor(86, 96, 112)
LIGHT_FILL = "F2F4F7"
BORDER = "DADCE0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, size=None, color=None, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 23, NAVY, 0, 6),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "OTC University Representative Capability Statement"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(header.runs[0], size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Overseas Tutorial Centre Ltd | overseasuk.com"
    style_run(footer.runs[0], size=9, color=MUTED)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    style_run(run, size=9, color=GOLD, bold=True)


def add_meta_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    style_run(r1, size=10.5, color=NAVY, bold=True)
    r2 = p.add_run(value)
    style_run(r2, size=10.5, color=MUTED)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        style_run(run, size=11, color=RGBColor(20, 32, 51))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        style_run(run, size=11, color=RGBColor(20, 32, 51))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        style_run(r, size=10.5, color=NAVY, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            style_run(r, size=10, color=RGBColor(42, 52, 66))
    doc.add_paragraph()
    return table


def build_docx():
    ASSETS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)

    add_kicker(doc, "Institutional Cooperation Brief")
    title = doc.add_paragraph(style="Title")
    title.add_run("Overseas Tutorial Centre Ltd")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("University Representative & Application Support Capability Statement")

    add_meta_row(doc, "Prepared for", "UK university admissions teams, international offices and partner managers")
    add_meta_row(doc, "Date", "19 May 2026")
    add_meta_row(doc, "Organisation", "Overseas Tutorial Centre Ltd, London, United Kingdom")
    add_meta_row(doc, "Company No.", "11060519")
    add_meta_row(doc, "Website", "https://overseasuk.com")
    add_meta_row(doc, "Contact", "office@overseasuk.com | +44 7947 991572")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "OTC supports Chinese and international students with UK university application preparation, "
        "advanced-entry evidence, bilingual document organisation, admissions communication and "
        "education-route coordination. OTC is seeking appropriate formal representative arrangements "
        "with UK universities where institutional policies, training, data rules and admissions processes "
        "can be followed properly."
    )
    style_run(r, size=11.5, color=RGBColor(20, 32, 51))

    doc.add_heading("1. Institutional Profile", level=1)
    add_bullets(
        doc,
        [
            "UK-based education consulting, tutoring, publishing and digital learning organisation.",
            "Public website: https://overseasuk.com, with dedicated UK university application, service-standard and institutional cooperation pages.",
            "Student audience includes Chinese university international programme students, international-school learners, pathway learners, undergraduate applicants and postgraduate applicants.",
            "Operational focus: careful document handling, realistic eligibility screening, admissions-facing communication and clear service boundaries.",
        ],
    )

    doc.add_heading("2. Application Support Scope", level=1)
    add_table(
        doc,
        ["Route", "OTC support", "Important boundary"],
        [
            [
                "Undergraduate / UCAS",
                "Course screening, document checklist, personal statement planning notes, references and deadline map.",
                "UCAS or university rules remain controlling.",
            ],
            [
                "Direct applications",
                "Admissions enquiry drafting, application route planning, file naming and offer-condition follow-up.",
                "Final admission decision remains with the university.",
            ],
            [
                "Year 2 / advanced entry",
                "Transcript review, module summaries, course descriptions, grading context and academic mapping notes.",
                "Advanced standing and credit are discretionary.",
            ],
            [
                "Postgraduate taught",
                "Course portfolio, CV/PS planning, evidence organisation, English-readiness and timing review.",
                "No guarantee of offer, scholarship or visa outcome.",
            ],
        ],
        [1.55, 3.25, 1.7],
    )

    doc.add_heading("3. China Programme and Advanced-Entry Capability", level=1)
    doc.add_paragraph(
        "OTC handles repeated enquiries from students studying in China-based international undergraduate programmes, "
        "including cases where the student has completed Year 1 or Year 2 modules and needs a UK receiving institution "
        "to advise whether Year 2, top-up, advanced standing or an alternative entry route can be considered."
    )
    add_bullets(
        doc,
        [
            "Student profile and target-course summary.",
            "Transcript, grading scale, English evidence and high-school record where requested.",
            "Course descriptions, module summaries, credit/contact-hour notes and timetable evidence where available.",
            "Module mapping against the intended UK course, with gaps flagged rather than hidden.",
            "Admissions enquiry wording that makes clear the university decides final entry level and suitability.",
        ],
    )

    doc.add_heading("4. Document Control and Compliance Standards", level=1)
    add_table(
        doc,
        ["Standard", "OTC position"],
        [
            ["Student consent", "OTC works from student instructions and records who may receive updates."],
            ["Document authenticity", "OTC does not create, alter or endorse false transcripts, certificates, references, scores or employment evidence."],
            ["Data protection", "Application files are used for the agreed purpose and shared only with relevant staff, advisers or institutions."],
            ["Academic integrity", "OTC may help students plan and express their own application materials, but does not complete assessed work or impersonate students."],
            ["Admissions boundaries", "OTC does not guarantee admission, scholarships, credit transfer, advanced standing, visa outcomes or appeal success."],
            ["Regulated advice", "Immigration, legal, tax or regulated professional advice is referred to appropriately qualified professionals where required."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("5. Partner Readiness", level=1)
    add_bullets(
        doc,
        [
            "OTC can follow university agent training, portal requirements, branding rules and reporting expectations where a formal arrangement is agreed.",
            "OTC can maintain enquiry logs, missing-document lists, deadline records and offer-condition tracking for institutional review.",
            "OTC can separate public application information from private commercial or representative terms.",
            "Unless a written agreement says otherwise, references to universities on OTC public pages are contextual information only and do not imply endorsement or official partnership.",
        ],
    )

    doc.add_heading("6. Cooperation Request", level=1)
    doc.add_paragraph(
        "OTC would welcome guidance on the university's official representative or agent onboarding process, "
        "including training requirements, application portal workflow, data-protection expectations, territory "
        "or student-category rules, and whether advanced-entry pre-checks may be handled through a defined channel."
    )
    add_numbered(
        doc,
        [
            "Confirm the university's agent or representative application process.",
            "Confirm whether OTC may submit or coordinate enquiries for China programme students needing advanced-entry review.",
            "Confirm required training, compliance checks, student-consent rules and admissions communication route.",
            "Agree how applications, document checks, offer conditions and follow-up communications should be recorded.",
        ],
    )

    doc.add_heading("Key Web Pages", level=1)
    add_bullets(
        doc,
        [
            "UK University Applications: https://overseasuk.com/university-applications/",
            "Application Service Standards: https://overseasuk.com/application-service-standards/",
            "Advanced Entry & China Programme Support: https://overseasuk.com/advanced-entry-china-programmes/",
            "University Agent & Institutional Cooperation: https://overseasuk.com/university-partnerships/",
        ],
    )

    doc.save(DOCX_OUT)
    return DOCX_OUT


def p(text, style):
    return Paragraph(text, style)


def bullet_list(items, style):
    return ListFlowable(
        [ListItem(Paragraph(item, style), leftIndent=14) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=16,
    )


def numbered_list(items, style):
    return ListFlowable(
        [ListItem(Paragraph(item, style), leftIndent=14) for item in items],
        bulletType="1",
        leftIndent=16,
    )


def pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(colors.white)
    canvas.rect(0, 0, letter[0], letter[1], stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor("#566070"))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(inch, 10.45 * inch, "OTC University Representative Capability Statement")
    canvas.drawRightString(7.5 * inch, 0.55 * inch, "Overseas Tutorial Centre Ltd | overseasuk.com")
    canvas.restoreState()


def build_pdf():
    ASSETS.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Kicker",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#B88A3B"),
            spaceAfter=6,
            uppercase=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="DocTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=27,
            textColor=colors.HexColor("#061A2F"),
            alignment=TA_LEFT,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="DocSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#566070"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor("#142033"),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            parent=styles["Body"],
            fontName="Helvetica",
            fontSize=9.8,
            leading=13,
            textColor=colors.HexColor("#566070"),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableText",
            parent=styles["Body"],
            fontSize=9,
            leading=12,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableHead",
            parent=styles["TableText"],
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#061A2F"),
        )
    )

    story = [
        p("Institutional Cooperation Brief".upper(), styles["Kicker"]),
        p("Overseas Tutorial Centre Ltd", styles["DocTitle"]),
        p("University Representative & Application Support Capability Statement", styles["DocSubtitle"]),
    ]
    meta_rows = [
        ("Prepared for", "UK university admissions teams, international offices and partner managers"),
        ("Date", "19 May 2026"),
        ("Organisation", "Overseas Tutorial Centre Ltd, London, United Kingdom"),
        ("Company No.", "11060519"),
        ("Website", "https://overseasuk.com"),
        ("Contact", "office@overseasuk.com | +44 7947 991572"),
    ]
    for label, value in meta_rows:
        story.append(p(f"<b>{label}:</b> {value}", styles["Meta"]))
    story.append(Spacer(1, 12))
    story.append(
        p(
            "OTC supports Chinese and international students with UK university application preparation, "
            "advanced-entry evidence, bilingual document organisation, admissions communication and "
            "education-route coordination. OTC is seeking appropriate formal representative arrangements "
            "with UK universities where institutional policies, training, data rules and admissions processes "
            "can be followed properly.",
            styles["Body"],
        )
    )

    story.append(p("1. Institutional Profile", styles["H1"]))
    story.append(
        bullet_list(
            [
                "UK-based education consulting, tutoring, publishing and digital learning organisation.",
                "Public website with dedicated UK university application, service-standard and institutional cooperation pages.",
                "Student audience includes Chinese university international programme students, international-school learners, pathway learners, undergraduate applicants and postgraduate applicants.",
                "Operational focus: careful document handling, realistic eligibility screening, admissions-facing communication and clear service boundaries.",
            ],
            styles["Body"],
        )
    )

    story.append(p("2. Application Support Scope", styles["H1"]))
    rows = [
        [p("<b>Route</b>", styles["TableHead"]), p("<b>OTC support</b>", styles["TableHead"]), p("<b>Important boundary</b>", styles["TableHead"])],
        [p("Undergraduate / UCAS", styles["TableText"]), p("Course screening, document checklist, personal statement planning notes, references and deadline map.", styles["TableText"]), p("UCAS or university rules remain controlling.", styles["TableText"])],
        [p("Direct applications", styles["TableText"]), p("Admissions enquiry drafting, application route planning, file naming and offer-condition follow-up.", styles["TableText"]), p("Final admission decision remains with the university.", styles["TableText"])],
        [p("Year 2 / advanced entry", styles["TableText"]), p("Transcript review, module summaries, course descriptions, grading context and academic mapping notes.", styles["TableText"]), p("Advanced standing and credit are discretionary.", styles["TableText"])],
        [p("Postgraduate taught", styles["TableText"]), p("Course portfolio, CV/PS planning, evidence organisation, English-readiness and timing review.", styles["TableText"]), p("No guarantee of offer, scholarship or visa outcome.", styles["TableText"])],
    ]
    t = Table(rows, colWidths=[1.55 * inch, 3.25 * inch, 1.7 * inch], repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(t)

    story.append(PageBreak())
    story.append(p("3. China Programme and Advanced-Entry Capability", styles["H1"]))
    story.append(
        p(
            "OTC handles repeated enquiries from students studying in China-based international undergraduate programmes, "
            "including cases where the student has completed Year 1 or Year 2 modules and needs a UK receiving institution "
            "to advise whether Year 2, top-up, advanced standing or an alternative entry route can be considered.",
            styles["Body"],
        )
    )
    story.append(
        bullet_list(
            [
                "Student profile and target-course summary.",
                "Transcript, grading scale, English evidence and high-school record where requested.",
                "Course descriptions, module summaries, credit/contact-hour notes and timetable evidence where available.",
                "Module mapping against the intended UK course, with gaps flagged rather than hidden.",
                "Admissions enquiry wording that makes clear the university decides final entry level and suitability.",
            ],
            styles["Body"],
        )
    )

    story.append(p("4. Document Control and Compliance Standards", styles["H1"]))
    rows = [[p("<b>Standard</b>", styles["TableHead"]), p("<b>OTC position</b>", styles["TableHead"])]]
    for standard, position in [
        ("Student consent", "OTC works from student instructions and records who may receive updates."),
        ("Document authenticity", "OTC does not create, alter or endorse false transcripts, certificates, references, scores or employment evidence."),
        ("Data protection", "Application files are used for the agreed purpose and shared only with relevant staff, advisers or institutions."),
        ("Academic integrity", "OTC may help students plan and express their own application materials, but does not complete assessed work or impersonate students."),
        ("Admissions boundaries", "OTC does not guarantee admission, scholarships, credit transfer, advanced standing, visa outcomes or appeal success."),
        ("Regulated advice", "Immigration, legal, tax or regulated professional advice is referred to appropriately qualified professionals where required."),
    ]:
        rows.append([p(standard, styles["TableText"]), p(position, styles["TableText"])])
    t2 = Table(rows, colWidths=[1.8 * inch, 4.7 * inch], repeatRows=1)
    t2.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(t2)

    story.append(p("5. Partner Readiness", styles["H1"]))
    story.append(
        bullet_list(
            [
                "OTC can follow university agent training, portal requirements, branding rules and reporting expectations where a formal arrangement is agreed.",
                "OTC can maintain enquiry logs, missing-document lists, deadline records and offer-condition tracking for institutional review.",
                "OTC can separate public application information from private commercial or representative terms.",
                "Unless a written agreement says otherwise, references to universities on OTC public pages are contextual information only and do not imply endorsement or official partnership.",
            ],
            styles["Body"],
        )
    )

    story.append(p("6. Cooperation Request", styles["H1"]))
    story.append(
        p(
            "OTC would welcome guidance on the university's official representative or agent onboarding process, "
            "including training requirements, application portal workflow, data-protection expectations, territory "
            "or student-category rules, and whether advanced-entry pre-checks may be handled through a defined channel.",
            styles["Body"],
        )
    )
    story.append(
        numbered_list(
            [
                "Confirm the university's agent or representative application process.",
                "Confirm whether OTC may submit or coordinate enquiries for China programme students needing advanced-entry review.",
                "Confirm required training, compliance checks, student-consent rules and admissions communication route.",
                "Agree how applications, document checks, offer conditions and follow-up communications should be recorded.",
            ],
            styles["Body"],
        )
    )

    story.append(p("Key Web Pages", styles["H1"]))
    story.append(
        bullet_list(
            [
                "UK University Applications: https://overseasuk.com/university-applications/",
                "Application Service Standards: https://overseasuk.com/application-service-standards/",
                "Advanced Entry & China Programme Support: https://overseasuk.com/advanced-entry-china-programmes/",
                "University Agent & Institutional Cooperation: https://overseasuk.com/university-partnerships/",
            ],
            styles["Body"],
        )
    )

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)
    return PDF_OUT


if __name__ == "__main__":
    print(f"Wrote {build_docx()}")
    print(f"Wrote {build_pdf()}")
